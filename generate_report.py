from io import BytesIO
from flask import Flask, request, send_file
import mcq_analyzer 
import docx
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from collections import defaultdict
import tempfile
app = Flask(__name__)


@app.route("/")
def index():
	return """<!DOCTYPE html>
<html>
<head>
	<title>Generate Word File</title>
</head>
<body>
	<form method="POST" action="/generate_word" enctype="multipart/form-data">
		<label for="file1">File 1:</label>
		<input type="file" id="file1" name="file1"><br><br>

		<label for="file2">File 2:</label>
		<input type="file" id="file2" name="file2"><br><br>

		<input type="submit" value="Generate">
	</form>
</body>
</html>"""
def read_file(file_name):
    with open(file_name, "r") as f:
        lines = f.readlines()
    return [line.strip() for line in lines]


def get_topic_scores(topic_dict):
    correct = sum(topic_dict.values())
    total = len(topic_dict)
    return correct, total

@app.route("/generate_word", methods=["POST"])
def generate_word():
    # Open the two files
    file1 = request.files["file1"]
    file2 = request.files["file2"] 
    answer_key = read_file(file1)
    student_answers = read_file(file2)

    reading_dict = defaultdict(lambda: defaultdict(int))
    writing_dict = defaultdict(lambda: defaultdict(int))
    section = None

    for i in range(len(answer_key)):
        if answer_key[i].startswith("Reading"):
            section = "reading"
            continue
        elif answer_key[i].startswith("Writing"):
            section = "writing"
            continue

        topic, answer = answer_key[i].split(",")
        if section == "reading":
            if answer == student_answers[i]:
                reading_dict[topic]["correct"] += 1
            else:
                reading_dict[topic]["incorrect"] += 1
        elif section == "writing":
            if answer == student_answers[i]:
                writing_dict[topic]["correct"] += 1
            else:
                writing_dict[topic]["incorrect"] += 1
	# Create a new Word document
    doc = docx.Document()

	    # Add the title
    title = doc.add_heading("Test Report", 0)
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    # Add the reading table
    doc.add_paragraph("Reading Section", style='Heading 1')
    reading_table = doc.add_table(rows=1, cols=3)
    reading_table.style = "Table Grid"
    hdr_cells = reading_table.rows[0].cells
    hdr_cells[0].text = "Topic"
    hdr_cells[1].text = "Correct"
    hdr_cells[2].text = "Incorrect"

    for topic, score_dict in reading_dict.items():
        correct, total = get_topic_scores(score_dict)
        percentage = (correct / total) * 100 if total > 0 else 0
        row_cells = reading_table.add_row().cells
        row_cells[0].text = topic
        row_cells[1].text = str(correct)
        row_cells[2].text = str(score_dict["incorrect"])

    # Add some more styling
    for row in reading_table.rows:
        for cell in row.cells:
            cell.width = docx.shared.Inches(2)
            paragraphs = cell.paragraphs
            for p in paragraphs:
                for run in p.runs:
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(11)

    # Add the writing table
    doc.add_paragraph("Writing Section", style='Heading 1')
    writing_table = doc.add_table(rows=1, cols=3)
    writing_table.style = "Table Grid"
    hdr_cells = writing_table.rows[0].cells
    hdr_cells[0].text = "Topic"
    hdr_cells[1].text = "Correct"
    hdr_cells[2].text = "Incorrect"

    for topic, score_dict in writing_dict.items():
        correct, total = get_topic_scores(score_dict)
        percentage = (correct / total) * 100 if total > 0 else 0
        row_cells = writing_table.add_row().cells
        row_cells[0].text = topic
        row_cells[1].text = str(correct)
        row_cells[2].text = str(score_dict["incorrect"])

    # Add some more styling
    for row in writing_table.rows:
        for cell in row.cells:
            cell.width = docx.shared.Inches(2)
            paragraphs = cell.paragraphs
            for p in paragraphs:
                for run in p.runs:
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(11)

    # Add a page break
    doc.add_page_break()

    # Add the overall score
    doc.add_heading("Overall Score", level=1)
    num_correct = sum([sum(d.values()) for d in reading_dict.values()] +
                      [sum(d.values()) for d in writing_dict.values()])
    num_questions = len(answer_key)
    overall_percentage = (num_correct / num_questions) * 100 if num_questions > 0 else 0
    doc.add_paragraph(
        f"You got {num_correct} out of {num_questions} questions correct, which is {overall_percentage:.2f}%.")

    # Save the document to the BytesIO object
    # doc.save(report_file)
    # return report_file

    # Add the conclusion
    conclusion = doc.add_paragraph("Well done! Keep up the good work!")
    conclusion.style.font.size = Pt(12)
    # Add a paragraph
    doc.add_paragraph("Here's a Google Drive link that includes material you can study from:")
    # Get the link text
    link_text = "https://drive.google.com/drive/folders/1V12ViiL2x_DuwJf1hqxRKQycf8cLiRrR?usp=share_link"
    link_paragraph = doc.add_paragraph(link_text)
    link_paragraph.style.font.size = Pt(12)
    
    doc_file = "temp.docx"
    
    # Save the document to a temporary file
    doc.save(doc_file)

	# Send the file to the user for download
    return send_file(doc_file, as_attachment=True)




if __name__ == "__main__":
    app.run(debug=False, host = '0.0.0.0')